from cgitb import text
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
# from docx.text.paragraph
import os

def valueEmbedder(template_path, output_path, data):
    doc = Document(template_path)
    
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if key in paragraph.text:
                for run in paragraph.runs:
                    if "[College Logo]" in key or "[Company Logo]" in key:
                        # print(key)
                        
                        print(os.getcwd())
                        paragraph.text=paragraph.text.replace(key," ")                       
                        
                        run = paragraph.add_run()
                        run.add_picture(f'{value}',width=Cm(4),height=Cm(4))
                        run.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        # print("Yes")
                        
                    
                    elif "[Number List Projects Purpose]" in key or "[Number List Projects Context]" in key:
                        paragraph.text=paragraph.text.replace(key," ")
                        for line in value:
                            paragraph.add_run(line+"\n")
                    
                    elif "[Signatory Signature]" in key or "[Company Head Signature]" in key:
                        paragraph.text=paragraph.text.replace(key," ")                       
                        run = paragraph.add_run(value['Name'])
                        run.font.name= value['Style']
                        # style = doc.styles['Normal']
                        # font = style.font
                        # font.name=value["Style"]
                        # signature = value['Name']
                        # paragraph.text = paragraph.text.replace(key, signature )
                        # paragraph.style=doc.styles["Normal"]
                        # run.font.name = value["Style"]
                        # paragraph.font.name=value["Style"]
                        
                        print(key, value["Name"])
                        # print(value['Style'])
                        # paragraph.add_run(value["Name"], style=value["Style"])
                        # paragraph.style = doc.styles["Normal"]
                        
                        # style = doc.styles["Normal"]
                        # font = style.font
                        # font.name=value["Style"]
                        # run.style = value["Style"]
                        pass
                        
                    
                    else:
                        paragraph.text = paragraph.text.replace(key, value)
                        pass
    
    doc.save(output_path)

if __name__=='__main__':
    os.chdir('backendwork/')
    # print(os.getcwd())
    # print(os.listdir())
    # print(os.listdir())
    data = {
        '[College Name]':'DSEU, Dwarka',
        '[Company Name]':'Blue Planet InfoSolutions',
        '[College Logo]':r'uploadedFiles\Q1NokiaAutoma2024-04-18_22-58-02-935825.jpg',
        '[Companay Logo]': r'uploadedFiles\companyLogo\blueplanet.png',
        '[Start Date]': '19-04-2024',
        '[Company Name Short]':'Blue Planet',
        '[College Address]':'Dwarka Sector 9, New Delhi - 110007',
        # '[Number List Projects Purpose]':["Purpose 1","Purpose 2","Purpose 3"],
        # "[Number List Projects Context]":["Context 1", "Context 2", "Context 3"],
        "[Signatory Signature]":{"Name":"Alpha Cafe","Style":"Vladimir Script"},
        "[Company Head Signature]":{"Name":"Beta Cafe","Style":"Brush Script MT"},
        "[Signatory]":"Alpha Cafe",
        "[Company Head]":"Beta Cafe",
        
    }
    # os.chdir('docswork/')
    
    template_path = r"docswork\templates\testtemplate.docx"
    output_path = r"docswork\results\result.docx"
    
    valueEmbedder(template_path, output_path, data)